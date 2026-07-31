"""수업 패키지 DoD 검사 (DESIGN.md §7 1단계 체크리스트).

사용: python tests/check_lesson.py <lesson.json> <outdir> [--official "성취기준 공식 원문"]

검사 항목:
  1. 렌더 무결성 — documents[] 전부 docx로 나왔고 U+FFFD 없음
  2. 성취기준 verbatim — shared.standard_text가 --official과 문자 단위로 일치
  3. 인용 1회 — 렌더된 교사 문서 전체에서 성취기준 원문이 정확히 한 번
  4. 시간 — phase_header minutes 합계 == shared.duration
  5. 3범주 — 지식·이해 / 과정·기능 / 가치·태도가 수업안에 모두 진술됨
"""
import argparse
import json
import os
import sys

from docx import Document

CATEGORIES = ["지식·이해", "과정·기능", "가치·태도"]


def docx_text(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def walk_blocks(blocks):
    for b in blocks or []:
        if not isinstance(b, dict):
            continue
        yield b
        for key in ("blocks", "left", "right"):
            yield from walk_blocks(b.get(key))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("lesson")
    ap.add_argument("outdir")
    ap.add_argument("--official", default=None,
                    help="학습맵 get_standard의 officialText (verbatim 대조용)")
    args = ap.parse_args()

    data = json.load(open(args.lesson, encoding="utf-8"))
    shared = data["shared"]
    failures = []
    texts = {}

    # 1. 렌더 무결성
    for doc in data["documents"]:
        path = os.path.join(args.outdir, doc["id"] + ".docx")
        if not os.path.exists(path):
            failures.append(f"[렌더] 누락: {path}")
            continue
        text = docx_text(path)
        texts[doc["id"]] = text
        if "�" in text:
            failures.append(f"[렌더] 깨진 문자: {path}")
        print(f"  렌더 ok: {doc['id']}.docx ({len(text)}자)")

    # 2. 성취기준 verbatim
    st = shared["standard_text"]
    if args.official is not None:
        if st == args.official:
            print("  verbatim ok: 학습맵 officialText와 문자 단위 일치")
        else:
            failures.append(f"[verbatim] 불일치\n    lesson: {st!r}\n    official: {args.official!r}")

    # 3. 인용 1회 (교사 문서 기준)
    for doc in data["documents"]:
        if doc.get("audience") != "teacher" or doc["id"] not in texts:
            continue
        n = texts[doc["id"]].count(st)
        if doc["id"] == "lesson_plan" and n != 1:
            failures.append(f"[인용] lesson_plan에 성취기준 원문 {n}회 (1회여야 함)")
        elif doc["id"] != "lesson_plan" and n > 1:
            failures.append(f"[인용] {doc['id']}에 성취기준 원문 {n}회")
    print("  인용 ok: 수업안에 성취기준 원문 1회")

    # 4. 시간
    minutes = [b.get("minutes", 0) for d in data["documents"]
               for s in d["sections"] for b in walk_blocks(s["blocks"])
               if b.get("type") == "phase_header"]
    total, duration = sum(minutes), shared["duration"]
    if total != duration:
        failures.append(f"[시간] 단계 합계 {total}분 != duration {duration}분 ({minutes})")
    else:
        print(f"  시간 ok: {' + '.join(map(str, minutes))} = {total}분")

    # 5. 3범주
    plan = texts.get("lesson_plan", "")
    missing = [c for c in CATEGORIES if c not in plan]
    if missing:
        failures.append(f"[3범주] 수업안에 없음: {', '.join(missing)}")
    else:
        print("  3범주 ok: 지식·이해 / 과정·기능 / 가치·태도 모두 진술됨")

    if failures:
        print("\nFAIL")
        for f in failures:
            print(" -", f)
        sys.exit(1)
    print("\nDoD 체크리스트 통과")


if __name__ == "__main__":
    main()
