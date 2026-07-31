"""렌더된 docx의 한글 무결성 검사: U+FFFD 없음 + 핵심 한글 문자열 존재."""
import glob
import sys

from docx import Document


def all_text(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def main(outdir):
    paths = sorted(glob.glob(f"{outdir}/*.docx"))
    assert paths, f"no docx in {outdir}"
    combined = ""
    for p in paths:
        text = all_text(p)
        assert "�" not in text, f"replacement character in {p}"
        combined += text
        print(f"ok: {p} ({len(text)} chars)")
    for needle in ["과학적으로 탐구할 수 있다", "탐구", "무엇이 보이나요", "설명하지 않고 제시"]:
        assert needle in combined, f"missing: {needle}"
    print("korean smoke test passed")


if __name__ == "__main__":
    main(sys.argv[1])
